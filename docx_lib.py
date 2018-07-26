
# -*- coding: utf-8 -*-
# encoding=utf8

from __future__ import unicode_literals

import os
import re
import sys
import json
import docx
import string
import traceback

from datetime						import datetime, date, time

from docx							import *
from docx							import Document
from docx.shared					import Cm, Inches
from docx.enum.text					import WD_ALIGN_PARAGRAPH
from docx.enum.text					import WD_LINE_SPACING
from docx.enum.table				import WD_TABLE_ALIGNMENT
from docx.enum.style				import WD_STYLE_TYPE, WD_STYLE
from docx.styles.style				import StyleFactory
from docx.text.run					import Font, Run
from docx.shared					import RGBColor
from docx.shared					import Pt
from docx.enum.style				import WD_STYLE
from docx.section					import *
from docx.oxml.section				import CT_SectPr

#----------------------------------------------------------------------------------------------------------------------------------------

reload( sys )
sys.setdefaultencoding( 'utf8' )


def make_docx () :
	
	#S O U R C E    O B J E C T S

	conditions	= {	}
	subtotals	=[]
	total = {}

	rtype		= 5
	error		= 0

	ttl_clr_prc_lst = []

	tbl_1_par_list 	= []
	tbl_1_par_list_cnr = []
	#----------------------------------------------------------------------------------------------------------------------------------------

	def get_cur_tmp( currency ):
		if currency == 'RUB' : result = u'''{:,.2f}р.'''
		if currency == 'RUR' : result = u'''{:,.2f}р.'''
		if currency == 'USD' : result = u'''${:,.2f}'''
		if currency == 'EUR' : result = u'''{:,.2f}€'''
		return result

	#----------------------------------------------------------------------------------------------------------------------------------------

	def add_tbl_line_srvs ( msg, prc_tpe ) :
		
		item_index 		= 1
		ttl_clr_prc_lst = []
		
		if prc_tpe == 'travel'	: adt_prc = total['Travel Price']
		if prc_tpe == 'total'	: adt_prc = total['Total Price']
		
		for subtotal in subtotals:
			for item in subtotal['items']:
				if item['Local Name']!= '':
						
					row_cls_1_1 = doc_tbl_1.add_row().cells
					
					tbl_str_1_1_x0 		= str(item_index)
					tbl_str_1_1_x1 		= str(item['partnumber'])
					tbl_str_1_1_x2 		= u'{ru_descr}'.format( ru_descr = item['Local Name'] )
					tbl_str_1_1_x3 		= str(float( item['Local Price'] ))
					tbl_str_1_1_x4 		= str(int( item['QTY'] ))
					tbl_str_1_1_x5 		= str(float(tbl_str_1_1_x3) * float(tbl_str_1_1_x4))
					
					ttl_clr_prc_lst.append(float(tbl_str_1_1_x5))
					
					row_cls_1_1[0].text 	= tbl_str_1_1_x0
					row_cls_1_1[1].text 	= tbl_str_1_1_x1
					row_cls_1_1[2].text 	= tbl_str_1_1_x2
					row_cls_1_1[3].text 	= crc_fmt.format( float( tbl_str_1_1_x3 ) )
					row_cls_1_1[4].text 	= tbl_str_1_1_x4
					row_cls_1_1[5].text 	= crc_fmt.format( float( tbl_str_1_1_x5 ) )
					
					tbl_1_par_1 	= row_cls_1_1[3].paragraphs
					tbl_1_par_2 	= row_cls_1_1[4].paragraphs
					tbl_1_par_3 	= row_cls_1_1[5].paragraphs
					
					for par in tbl_1_par_1: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_3: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_2: 	tbl_1_par_list_cnr.append(par)
					
					item_index = item_index + 1
					
				if item['Local Name'] == '':
					
					row_cls_1_1 = doc_tbl_1.add_row().cells
					
					tbl_str_1_1_x0 		= str(item_index)
					tbl_str_1_1_x1 		= str(item['partnumber'])
					tbl_str_1_1_x2 		= u'{en_descr}'.format( en_descr = item['Description'] )
					tbl_str_1_1_x3 		= str(float( item['Local Price'] ))
					tbl_str_1_1_x4 		= str(int( item['QTY'] ))
					tbl_str_1_1_x5 		= str(float(tbl_str_1_1_x3) * float(tbl_str_1_1_x4))
					
					ttl_clr_prc_lst.append(float(tbl_str_1_1_x5))
					
					row_cls_1_1[0].text 	= tbl_str_1_1_x0
					row_cls_1_1[1].text 	= tbl_str_1_1_x1
					row_cls_1_1[2].text 	= tbl_str_1_1_x2
					row_cls_1_1[3].text 	= crc_fmt.format( float( tbl_str_1_1_x3 ) )
					row_cls_1_1[4].text 	= tbl_str_1_1_x4
					row_cls_1_1[5].text 	= crc_fmt.format( float( tbl_str_1_1_x5 ) )
					
					tbl_1_par_1 	= row_cls_1_1[3].paragraphs
					tbl_1_par_2 	= row_cls_1_1[4].paragraphs
					tbl_1_par_3 	= row_cls_1_1[5].paragraphs
					
					for par in tbl_1_par_1: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_3: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_2: 	tbl_1_par_list_cnr.append(par) 
					
					item_index = item_index + 1
		
		row_cls_1_2 		= doc_tbl_1.add_row().cells
		
		tbl_str_1_2_x2 		= str( msg )
		tbl_str_1_2_x3 		= float( adt_prc )
		
		ttl_clr_prc_lst.append(float(tbl_str_1_2_x3))
		
		row_cls_1_2[0].text 	= str(item_index)
		row_cls_1_2[1].text 	= 'srv'
		
		try:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = u''.join( conditions['Work location']  ) )
		except UnicodeDecodeError:	row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = conditions['Work location'])
		except:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = 'Work location' )
		
		row_cls_1_2[3].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )
		row_cls_1_2[4].text 	= '1'
		row_cls_1_2[5].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )	
		
		tbl_1_par_12 	= row_cls_1_2[3].paragraphs
		tbl_1_par_22 	= row_cls_1_2[4].paragraphs
		tbl_1_par_32 	= row_cls_1_2[5].paragraphs
		
		for par in tbl_1_par_12: 	tbl_1_par_list.append(par)
		for par in tbl_1_par_32: 	tbl_1_par_list.append(par)
		for par in tbl_1_par_22: 	tbl_1_par_list_cnr.append(par)
		
		item_index = item_index + 1
		
		return item_index, ttl_clr_prc_lst
				
	# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
	# M A K E   D O C X   C O M M E R C I A L   O F F E R
	# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

	try:
		document				= Document()
		crc_fmt					= get_cur_tmp( conditions['Calculation currency'] )

		#json_string				= result


		pic_par_1				= document.add_paragraph()
		run_pic_1				= pic_par_1.add_run()
		pic_obj_1				= run_pic_1.add_picture( '{path}1.png', height = Inches( 1.40 ) )

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# M A K E   T I T L E   P A G E
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		ctr_num					= u'''ХХХХХХХХ'''								# contract	- number
		ctr_dte					= str(date.today())#jobcalc - get_date_string()	# contract	- creation date

		cmr_dsc					= u'''ООО "Company"'''							# customer	- company name
		cmr_nme					= u'''Алексей Алексеевич'''						# customer	- contract person
		cmr_phn					= u'''7-499-000-11-22'''						# customer	- contract person phone
		cmr_eml					= u'''luxury.escort@luxury.escort.cz'''			# customer	- contract person e-mail

		slr_pos					= u'''Менеджер'''			# se seller	- position
		slr_nme					= u'''Лапавекпне Н.'''							# seller	- name
		slr_phn					= u'''7-495-777-99-90'''						# seller	- phone
		slr_eml					= u'''Nikita.Lapatin@scle.com'''				# seller	- e-mail

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		#C O N T R A C T   I N F O R M A T I O N
		ctr_num					= '589659765'									# contract	- number
		ctr_dte					= str(date.today())#jobcalc - get_date_string()	# contract	- creation date
		ctr_mod					= '488-2'										# contract	- qid modification ex: 488-2

		cmr_dsc					= 'unassigned_customer_organisation'			# customer	- company name
		cmr_nme					= 'unassigned_customer_name'					# customer	- contract person
		cmr_phn					= 'unassigned_customer_phone'					# customer	- contract person phone
		cmr_eml					= 'unassigned_customer_email'					# customer	- contract person e-mail

		cmr_del_adr				= 'unassigned_customer_delivery_address'		# customer	- delivery address
		cmr_del_phn				= 'unassigned_customer_delivery_phone'			# customer	- delivery phone
		cmr_del_eml				= 'unassigned_customer_delivery_email'			# customer	- delivery email

		slr_pos					= 'unassigned_seller_position'					# se seller	- position
		slr_nme					= 'unassigned_seller_name'						# se seller	- name
		slr_phn					= 'unassigned_seller_phone'						# se seller	- phone
		slr_eml					= 'unassigned_seller_email'						# se seller	- e-mail

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		ttl_str_0				= u''''''
		ttl_str_1				= u'''КП №{ctr_num}'''.format( ctr_num = ctr_num )
		ttl_str_2				= u''' ''' 
		ttl_str_3				= u'''Заказчик: {cmr_dsc}\rКонтактное лицо: {cmr_nme}\rТел.: {cmr_phn}\rE-mail: {cmr_eml}\rДата: {ctr_dte}'''.format( cmr_dsc = cmr_dsc, cmr_nme = cmr_nme, cmr_phn = cmr_phn, cmr_eml = cmr_eml, ctr_dte = ctr_dte)
		ttl_str_4				= u''''''
		ttl_str_5				= u'''С уважением,\r{slr_pos}'''.format( slr_pos = slr_pos )
		ttl_str_6				= u'''{slr_nme}\r{slr_phn}\r{slr_eml}'''.format( slr_nme = slr_nme, slr_phn = slr_phn, slr_eml = slr_eml,  )
		ttl_str_7				= u'''Настоящее технико-коммерческое предложение (план) не может рассматриваться как оферта'''
		ttl_emp_0				= document.add_paragraph()					# Blank line
		ttl_emp_1				= document.add_paragraph()					# Blank line
		ttl_par_1	 			= document.add_paragraph()					# Commercial Offer ... #Number ...
		ttl_par_2				= document.add_paragraph()					# Blank line
		ttl_emp_2				= document.add_paragraph()					# Blank line
		ttl_emp_3				= document.add_paragraph()					# Blank line
		ttl_par_3 				= document.add_paragraph()					# Customer ... Contact Person ... Contact Phone ... Contact e-mail ... Date ...
		ttl_emp_4				= document.add_paragraph()					# Blank line
		ttl_par_4 				= document.add_paragraph()					# Company SE Offer to you ...
		ttl_emp_5				= document.add_paragraph()					# Blank line
		ttl_par_5 				= document.add_paragraph()					# Kind regards ... Seller Position 
		ttl_par_6 				= document.add_paragraph()					# Seller Name ... Seller Phone ...
		ttl_emp_6				= document.add_paragraph()					# Blank line
		ttl_par_7				= document.add_paragraph()					# This technical commercial offer cant be ...

		run_ttl_1				= ttl_par_1.add_run( ttl_str_1 )			# Service contract ...
		run_ttl_2 				= ttl_par_2.add_run( ttl_str_2 )			# Technical Commercial Offer ... #Opti ...
		run_ttl_3 				= ttl_par_3.add_run( ttl_str_3 )			# Customer ... Contact Person ... Contact Phone ... Contact e-mail ...
		run_ttl_4 				= ttl_par_4.add_run( ttl_str_4 )			# Company SE Offer to you ...
		run_ttl_5 				= ttl_par_5.add_run( ttl_str_5 )			# Kind regards ... Seller Position ...
		run_ttl_6 				= ttl_par_6.add_run( ttl_str_6 )			# Seller Name ... Seller Phone ...
		run_ttl_7 				= ttl_par_7.add_run( ttl_str_7 )			# This technical commercial offer cant be ...

		run_emp_0				= ttl_emp_0.add_run( ttl_str_0 )
		run_emp_1				= ttl_emp_1.add_run( ttl_str_0 )
		run_emp_2				= ttl_emp_2.add_run( ttl_str_0 )
		run_emp_3				= ttl_emp_3.add_run( ttl_str_0 )
		run_emp_4				= ttl_emp_4.add_run( ttl_str_0 )
		run_emp_5				= ttl_emp_5.add_run( ttl_str_0 )
		run_emp_6				= ttl_emp_6.add_run( ttl_str_0 )

		document.add_page_break()

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# P A R T   1
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		hed_str_1				= u'''1. ОБЩИЕ ПОЛОЖЕНИЯ И СРОКИ ДЕЙСТВИЯ'''


		doc_hed_1 				= document.add_heading( '', 0 )
		doc_par_10				= document.add_paragraph()
		doc_par_11				= document.add_paragraph()


		run_hed_1				= doc_hed_1.add_run( hed_str_1 )
		run_par_10 				= doc_par_10.add_run( par_str_10 )
		run_par_11 				= doc_par_11.add_run( par_str_11 )

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# P A R T   2
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		hed_str_2				= u'''2. ПОСТАВКА, ОСУЩЕСТВЛЯЕМАЯ ОAO "..."'''
		par_str_2				= u'''ОAO "..." предлагает в данном технико-коммерческом предложении следующие материалы:'''

		doc_hed_2 				= document.add_heading( '', 0 )
		doc_par_2				= document.add_paragraph()

		run_hed_2 				= doc_hed_2.add_run( hed_str_2 )
		run_par_2				= doc_par_2.add_run( par_str_2 )

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# M A K E  T A B L E   1
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		tbl_str_1_00				=u'''№'''
		tbl_str_1_01				=u'''Артикул'''
		tbl_str_1_02				=u'''Описание'''
		tbl_str_1_03				=u'''Цена за ед.\nв руб.'''
		tbl_str_1_04				=u'''Кол-во'''
		tbl_str_1_05				=u'''Стоимость в руб.\nбез НДС'''

		doc_tbl_1					= document.add_table( rows = 1, cols = 6, style = 'Table Grid' )

		doc_tbl_1.cell(0,0).text 	= tbl_str_1_00
		doc_tbl_1.cell(0,1).text 	= tbl_str_1_01
		doc_tbl_1.cell(0,2).text 	= tbl_str_1_02
		doc_tbl_1.cell(0,3).text 	= tbl_str_1_03
		doc_tbl_1.cell(0,4).text 	= tbl_str_1_04
		doc_tbl_1.cell(0,5).text 	= tbl_str_1_05

		#tbl_1_par_list 		= []
		#tbl_1_par_list_cnr 	= []
		item_index 			= 1

		if conditions['Request type'] == 'Services':

		#	CHECK SKU TYPES

			_abstract	= 0
			_service	= 0
			_device		= 0
			_both		= 0
			_contract	= 0

			for subtotal in subtotals:
				for item in subtotal['items']:
					if item['sku_type'] == 'abstract'	: _abstract	= _abstract	+ 1
					if item['sku_type'] == 'service'	: _service	= _service	+ 1
					if item['sku_type'] == 'device'		: _device	= _device	+ 1
					if item['sku_type'] == 'both'		: _both		= _both		+ 1
					if item['sku_type'] == 'contract'	: _contract	= _contract	+ 1

			#	ONLY SERVICES
			if 		( ( _abstract == 0 ) and ( _service != 0 ) and ( _device == 0 ) and ( _both == 0 ) and ( _contract == 0 ) ): item_index, ttl_clr_prc_lst = add_tbl_line_srvs(u'Региональная надбавка', 'travel')
				
			#	ONLY BOTH
			elif 	( ( _abstract == 0 ) and ( _service == 0 ) and ( _device == 0 ) and ( _both != 0 ) and ( _contract == 0 ) ): item_index, ttl_clr_prc_lst = add_tbl_line_srvs(u'Услуги по установке/замене комплектующих', 'total')

			#	MIX - SERVICES & BOTH
			elif 	( ( _abstract == 0 ) and ( _service != 0 ) and ( _device == 0 ) and ( _both != 0 ) and ( _contract == 0 ) ): item_index, ttl_clr_prc_lst = add_tbl_line_srvs(u'Региональная надбавка', 'travel')

			#	OTHER CASES
			else: item_index, ttl_clr_prc_lst = add_tbl_line_srvs(u'Услуги по установке/замене комплектующих', 'total')

		if conditions['Request type'] == 'Diagnostics':
			for subtotal in subtotals:
				for item in subtotal['items']:
							
					row_cls_1_1 = doc_tbl_1.add_row().cells
					
					tbl_str_1_1_x0 		= str(item_index)
					tbl_str_1_1_x1 		= str(item['partnumber'])
					tbl_str_1_1_x2 		= u'Диагностика оборудования ({en_descr})'.format( en_descr = item['Description'] ) 
					tbl_str_1_1_x3 		= str( total['work Price'] )
					tbl_str_1_1_x4 		= str(int( item['QTY'] ))
					tbl_str_1_1_x5 		= str(float(tbl_str_1_1_x3) * float(tbl_str_1_1_x4))
					
					ttl_clr_prc_lst.append(float(tbl_str_1_1_x5)) #price list for final price
					
					row_cls_1_1[0].text 	= tbl_str_1_1_x0
					row_cls_1_1[1].text 	= tbl_str_1_1_x1
					row_cls_1_1[2].text 	= tbl_str_1_1_x2
					row_cls_1_1[3].text 	= crc_fmt.format( float( tbl_str_1_1_x3 ) )
					row_cls_1_1[4].text 	= tbl_str_1_1_x4
					row_cls_1_1[5].text 	= crc_fmt.format( float( tbl_str_1_1_x5 ) )
					
					tbl_1_par_1 	= row_cls_1_1[3].paragraphs
					tbl_1_par_2 	= row_cls_1_1[4].paragraphs
					tbl_1_par_3 	= row_cls_1_1[5].paragraphs
					
					for par in tbl_1_par_1: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_3: 	tbl_1_par_list.append(par)
					for par in tbl_1_par_2: 	tbl_1_par_list_cnr.append(par)
					
					item_index = item_index + 1
						
			row_cls_1_2 		= doc_tbl_1.add_row().cells
			
			tbl_str_1_2_x2 		= u'Региональная надбавка'
			tbl_str_1_2_x3 		= float( total['Travel Price'] )
			
			ttl_clr_prc_lst.append(float(tbl_str_1_2_x3))
			
			row_cls_1_2[0].text 	= str(item_index)
			row_cls_1_2[1].text 	= 'srv'
			
			try:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = u''.join( conditions['Work location']  ) )
			except UnicodeDecodeError:	row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = conditions['Work location'])
			except:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = 'Work location' )
			
			row_cls_1_2[3].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )
			row_cls_1_2[4].text 	= '1'
			row_cls_1_2[5].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )	
			
			tbl_1_par_12 	= row_cls_1_2[3].paragraphs
			tbl_1_par_22 	= row_cls_1_2[4].paragraphs
			tbl_1_par_32 	= row_cls_1_2[5].paragraphs
			
			for par in tbl_1_par_12: 	tbl_1_par_list.append(par)
			for par in tbl_1_par_32: 	tbl_1_par_list.append(par)
			for par in tbl_1_par_22: 	tbl_1_par_list_cnr.append(par)
			
		if ( (conditions['Request type'] == 'Simple') or (conditions['Request type'] == 'Solutions') or (conditions['Request type'] == 'Batteries') ):
			if str( subtotals[0]['Equipment Family'] ) != 'srv':
				for subtotal in subtotals:
					for item in subtotal['items']:
								
						row_cls_1_1 = doc_tbl_1.add_row().cells
						
						tbl_str_1_1_x0 			= str(item_index)
						tbl_str_1_1_x1 			= str(item['partnumber'])
						tbl_str_1_1_x2 			= u'{ru_descr} ({en_descr})'.format( ru_descr = item['Local Name'], en_descr = item['Description'] )
						tbl_str_1_1_x3 			= str( item['Local Price'] )
						tbl_str_1_1_x4 			= str(int( item['QTY'] ))
						tbl_str_1_1_x5 			= str(float(tbl_str_1_1_x3) * float(tbl_str_1_1_x4))
						
						ttl_clr_prc_lst.append(float(tbl_str_1_1_x5))
						
						row_cls_1_1[0].text 	= tbl_str_1_1_x0
						row_cls_1_1[1].text 	= tbl_str_1_1_x1
						row_cls_1_1[2].text 	= tbl_str_1_1_x2
						row_cls_1_1[3].text 	= crc_fmt.format( float( tbl_str_1_1_x3 ) )
						row_cls_1_1[4].text 	= tbl_str_1_1_x4
						row_cls_1_1[5].text 	= crc_fmt.format( float( tbl_str_1_1_x5 ) )
						
						tbl_1_par_1 			= row_cls_1_1[3].paragraphs
						tbl_1_par_2 			= row_cls_1_1[4].paragraphs
						tbl_1_par_3 			= row_cls_1_1[5].paragraphs

						for par in tbl_1_par_1: 	tbl_1_par_list.append(par)
						for par in tbl_1_par_3: 	tbl_1_par_list.append(par)
						for par in tbl_1_par_2: 	tbl_1_par_list_cnr.append(par)
						
						item_index = item_index + 1
						
			row_cls_1_2 		= doc_tbl_1.add_row().cells
			
			tbl_str_1_2_x2 		= u'Услуги по установке/замене комплектующих'
			tbl_str_1_2_x3 		= float( total['Total Price'] )
			
			ttl_clr_prc_lst.append(float(tbl_str_1_2_x3))
			
			row_cls_1_2[0].text 	= str(item_index)
			row_cls_1_2[1].text 	= 'srv'
			
			try:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = u''.join( conditions['Work location']  ) )
			except UnicodeDecodeError:	row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = conditions['Work location'])
			except:						row_cls_1_2[2].text 	= '{msg} ({target})'.format( msg = tbl_str_1_2_x2, target = 'Work location' )
			
			row_cls_1_2[3].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )
			row_cls_1_2[4].text 	= '1'
			row_cls_1_2[5].text 	= crc_fmt.format( float( tbl_str_1_2_x3 ) )	
		
			tbl_1_par_12 			= row_cls_1_2[3].paragraphs
			tbl_1_par_22 			= row_cls_1_2[4].paragraphs
			tbl_1_par_32 			= row_cls_1_2[5].paragraphs

			for par in tbl_1_par_12: 	tbl_1_par_list.append(par)
			for par in tbl_1_par_32: 	tbl_1_par_list.append(par)
			for par in tbl_1_par_22: 	tbl_1_par_list_cnr.append(par)
			
			
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# P A R T   3
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		hed_str_4				= u'''Примечания'''
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# M A K E  T A B L E   2
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		tbl_str_2_00			= u'''Итого, руб без НДС'''
		tbl_str_2_10			= u'''НДС (18%)'''
		tbl_str_2_20			= u'''Итого, руб с НДС (18%)'''

		ttl_clr_prc 			= 0
		ttl_vat_prc				= 0

		for clr_prc in ttl_clr_prc_lst : 
			ttl_clr_prc 		= clr_prc + ttl_clr_prc
			
		ttl_vat_prc 			= ttl_clr_prc * 1.18

		tbl_str_2_01 			= str( ttl_clr_prc )
		tbl_str_2_11 			= str( float( ttl_vat_prc ) - float( ttl_clr_prc ) )
		tbl_str_2_21			= str( float( ttl_vat_prc ) )

		tbl_str_2_01			= crc_fmt.format( float( tbl_str_2_01 ) )
		tbl_str_2_11 			= crc_fmt.format( float( tbl_str_2_11 ) )
		tbl_str_2_21 			= crc_fmt.format( float( tbl_str_2_21 ) )

		doc_tbl_2					= document.add_table( rows = 3, cols = 2, style = 'Table Grid' )
		
		doc_tbl_2.cell(0,0).text 	= tbl_str_2_00
			
		doc_tbl_2.cell(1,0).text 	= tbl_str_2_10
		doc_tbl_2.cell(2,0).text 	= tbl_str_2_20
			
		doc_tbl_2.cell(0,1).text 	= tbl_str_2_01
		doc_tbl_2.cell(1,1).text 	= tbl_str_2_11
		doc_tbl_2.cell(2,1).text 	= tbl_str_2_21
		
		tbl_2_par_list = []
		
		for row in range(0,3):
			tbl_2_par 	= doc_tbl_2.cell(row,1).paragraphs
			
			for par in tbl_2_par:
				tbl_2_par_list.append(par)
					
		doc_hed_4 					= document.add_heading( '', 1 )
		doc_par_41					= document.add_paragraph()
		doc_par_42					= document.add_paragraph()
		doc_par_43					= document.add_paragraph()
		doc_par_44					= document.add_paragraph()
		doc_par_45					= document.add_paragraph()

		run_hed_4 					= doc_hed_4.add_run( hed_str_4 )
		run_par_41					= doc_par_41.add_run( par_str_41 )
		run_par_42					= doc_par_42.add_run( par_str_42 )
		run_par_43					= doc_par_43.add_run( par_str_43 )
		run_par_44					= doc_par_44.add_run( par_str_44 )
		run_par_45					= doc_par_45.add_run( par_str_45 )

		document.add_page_break()

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# A N N E X   1
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		anx1_hed_0			= document.add_heading( '', 0 )
		anx1_par_0			= document.add_paragraph()
		anx1_par_00			= document.add_paragraph()
		anx1_par_01			= document.add_paragraph()
		anx1_par_02			= document.add_paragraph()
		anx1_par_03			= document.add_paragraph()
		anx1_par_04			= document.add_paragraph()
		anx1_hed_1			= document.add_heading( '', 1 )
		anx1_par_11			= document.add_paragraph()
		anx1_par_12			= document.add_paragraph()
		anx1_hed_2			= document.add_heading( '', 1 )
		anx1_par_20			= document.add_paragraph()
		anx1_par_21			= document.add_paragraph()
		anx1_par_22			= document.add_paragraph()
		anx1_par_23			= document.add_paragraph()
		anx1_par_24			= document.add_paragraph()
		anx1_par_25			= document.add_paragraph()
		anx1_par_26			= document.add_paragraph()
		anx1_par_27			= document.add_paragraph()
		anx1_par_28			= document.add_paragraph()
		anx1_par_29			= document.add_paragraph()
		anx1_hed_3			= document.add_heading( '', 1 )
		anx1_par_30			= document.add_paragraph()
		anx1_par_31			= document.add_paragraph()
		anx1_par_32			= document.add_paragraph()
		anx1_par_33			= document.add_paragraph()
		anx1_hed_4			= document.add_heading( '', 1 )
		anx1_par_41			= document.add_paragraph()
		anx1_par_42			= document.add_paragraph()
		anx1_par_43			= document.add_paragraph()
		anx1_hed_5			= document.add_heading( '', 1 )
		anx1_par_51			= document.add_paragraph()
		anx1_par_52			= document.add_paragraph()
		anx1_par_53			= document.add_paragraph()
		anx1_hed_6			= document.add_heading( '', 1 )
		anx1_par_61			= document.add_paragraph()
		anx1_par_62			= document.add_paragraph()
		anx1_par_63			= document.add_paragraph()
		anx1_par_64			= document.add_paragraph()
		anx1_par_65			= document.add_paragraph()
		anx1_hed_7			= document.add_heading( '', 1 )
		anx1_par_71			= document.add_paragraph()
		anx1_par_72			= document.add_paragraph()
		anx1_par_73			= document.add_paragraph()
		anx1_hed_8			= document.add_heading( '', 1 )
		anx1_par_81			= document.add_paragraph()
		anx1_par_82			= document.add_paragraph()

		anx1_run_hed_0		= anx1_hed_0.add_run(anx1_hed_str_0)
		anx1_run_hed_1		= anx1_hed_1.add_run(anx1_hed_str_1)
		anx1_run_hed_2		= anx1_hed_2.add_run(anx1_hed_str_2)
		anx1_run_hed_3		= anx1_hed_3.add_run(anx1_hed_str_3)
		anx1_run_hed_4		= anx1_hed_4.add_run(anx1_hed_str_4)
		anx1_run_hed_5		= anx1_hed_5.add_run(anx1_hed_str_5)
		anx1_run_hed_6		= anx1_hed_6.add_run(anx1_hed_str_6)
		anx1_run_hed_7		= anx1_hed_7.add_run(anx1_hed_str_7)
		anx1_run_hed_8		= anx1_hed_8.add_run(anx1_hed_str_8)

		anx1_run_nme_1		= anx1_par_01.add_run(anx1_run_nme_1)
		anx1_run_nme_2		= anx1_par_02.add_run(anx1_run_nme_2)
		anx1_run_nme_3		= anx1_par_03.add_run(anx1_run_nme_3)
		anx1_run_nme_4		= anx1_par_04.add_run(anx1_run_nme_4)

		anx1_run_par_00		= anx1_par_00.add_run(anx1_par_str_00)
		anx1_run_par_01		= anx1_par_01.add_run(anx1_par_str_01)
		anx1_run_par_02		= anx1_par_02.add_run(anx1_par_str_02)
		anx1_run_par_03		= anx1_par_03.add_run(anx1_par_str_03)
		anx1_run_par_04		= anx1_par_04.add_run(anx1_par_str_04)

		anx1_run_par_11		= anx1_par_11.add_run(anx1_par_str_11)
		anx1_run_par_12		= anx1_par_12.add_run(anx1_par_str_12)
		anx1_run_par_20		= anx1_par_20.add_run(anx1_par_str_20)
		anx1_run_par_21		= anx1_par_21.add_run(anx1_par_str_21)
		anx1_run_par_22		= anx1_par_22.add_run(anx1_par_str_22)
		anx1_run_par_23		= anx1_par_23.add_run(anx1_par_str_23)
		anx1_run_par_24		= anx1_par_24.add_run(anx1_par_str_24)
		anx1_run_par_25		= anx1_par_25.add_run(anx1_par_str_25)
		anx1_run_par_26		= anx1_par_26.add_run(anx1_par_str_26)
		anx1_run_par_27		= anx1_par_27.add_run(anx1_par_str_27)
		anx1_run_par_28		= anx1_par_28.add_run(anx1_par_str_28)
		anx1_run_par_29		= anx1_par_29.add_run(anx1_par_str_29)
		anx1_run_par_30		= anx1_par_30.add_run(anx1_par_str_30)
		anx1_run_par_31		= anx1_par_31.add_run(anx1_par_str_31)
		anx1_run_par_32		= anx1_par_32.add_run(anx1_par_str_32)
		anx1_run_par_33		= anx1_par_33.add_run(anx1_par_str_33)
		anx1_run_par_41		= anx1_par_41.add_run(anx1_par_str_41)
		anx1_run_par_42		= anx1_par_42.add_run(anx1_par_str_42)
		anx1_run_par_43		= anx1_par_43.add_run(anx1_par_str_43)
		anx1_run_par_51		= anx1_par_51.add_run(anx1_par_str_51)
		anx1_run_par_52		= anx1_par_52.add_run(anx1_par_str_52)
		anx1_run_par_53		= anx1_par_53.add_run(anx1_par_str_53)
		anx1_run_par_61		= anx1_par_61.add_run(anx1_par_str_61)
		anx1_run_par_62		= anx1_par_62.add_run(anx1_par_str_62)
		anx1_run_par_63		= anx1_par_63.add_run(anx1_par_str_63)
		anx1_run_par_64		= anx1_par_64.add_run(anx1_par_str_64)
		anx1_run_par_65		= anx1_par_65.add_run(anx1_par_str_65)
		anx1_run_par_71		= anx1_par_71.add_run(anx1_par_str_71)
		anx1_run_par_72		= anx1_par_72.add_run(anx1_par_str_72)
		anx1_run_par_73		= anx1_par_73.add_run(anx1_par_str_73)
		anx1_run_par_81		= anx1_par_81.add_run(anx1_par_str_81)
		anx1_run_par_82		= anx1_par_82.add_run(anx1_par_str_82)

		document.add_page_break()

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# A N N E X   2
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		anx2_str_0			= u''''''
		anx2_hed_str_0 		= u'''10'''
		anx2_hed_str_1		= u'''1'''
		anx2_hed_str_2		= u'''2'''
		anx2_hed_str_3		= u'''3'''
		anx2_hed_str_4		= u'''4'''
		anx2_hed_str_5		= u'''5'''
		anx2_hed_str_6		= u'''6'''
		anx2_hed_str_7		= u'''7'''

		anx2_par_str_00		= u'''Следующие термины имеют в данном документе следующее значение:'''
		anx2_run_nme_1		= u''''''
		anx2_run_nme_2		= u'''Товар'''
		anx2_run_nme_3		= u'''Клиент'''
		anx2_run_nme_4		= u'''Заказ'''

		anx2_par_str_01		= ''
		anx2_par_str_02		= ''
		anx2_par_str_03		= ''

		anx2_hed_0			= document.add_heading( '', 0 )
		anx2_par_0			= document.add_paragraph()
		anx2_par_00			= document.add_paragraph()
		anx2_par_01			= document.add_paragraph()
		anx2_par_02			= document.add_paragraph()
		anx2_par_03			= document.add_paragraph()
		anx2_par_04			= document.add_paragraph()
		anx2_hed_1			= document.add_heading( '', 1 )
		anx2_par_11			= document.add_paragraph('' )
		anx2_par_12			= document.add_paragraph()
		anx2_hed_2			= document.add_heading( '', 1 )
		anx2_par_21			= document.add_paragraph()
		anx2_par_22			= document.add_paragraph()
		anx2_par_23			= document.add_paragraph()
		anx2_par_24			= document.add_paragraph()
		anx2_par_25			= document.add_paragraph()
		anx2_par_26			= document.add_paragraph()
		anx2_hed_3			= document.add_heading( '', 1 )
		anx2_par_31			= document.add_paragraph()
		anx2_par_32			= document.add_paragraph()
		anx2_par_33			= document.add_paragraph()
		anx2_par_34			= document.add_paragraph()
		anx2_par_35			= document.add_paragraph()
		anx2_hed_4			= document.add_heading( '', 1 )
		anx2_par_41			= document.add_paragraph()
		anx2_par_42			= document.add_paragraph()
		anx2_par_43			= document.add_paragraph()
		anx2_par_44			= document.add_paragraph()
		anx2_hed_5			= document.add_heading( '', 1 )
		anx2_par_51			= document.add_paragraph()
		anx2_par_52			= document.add_paragraph()
		anx2_par_53			= document.add_paragraph()
		anx2_par_54			= document.add_paragraph()
		anx2_par_55			= document.add_paragraph()
		anx2_par_56			= document.add_paragraph()
		anx2_hed_6			= document.add_heading( '', 1 )
		anx2_par_61			= document.add_paragraph()
		anx2_par_62			= document.add_paragraph()
		anx2_par_63			= document.add_paragraph()
		anx2_par_64			= document.add_paragraph()
		anx2_hed_7			= document.add_heading( '', 1 )
		anx2_par_71			= document.add_paragraph()
		anx2_par_72			= document.add_paragraph()

		anx2_run_hed_0		= anx2_hed_0.add_run(anx2_hed_str_0)
		anx2_run_hed_1		= anx2_hed_1.add_run(anx2_hed_str_1)
		anx2_run_hed_2		= anx2_hed_2.add_run(anx2_hed_str_2)
		anx2_run_hed_3		= anx2_hed_3.add_run(anx2_hed_str_3)
		anx2_run_hed_4		= anx2_hed_4.add_run(anx2_hed_str_4)
		anx2_run_hed_5		= anx2_hed_5.add_run(anx2_hed_str_5)
		anx2_run_hed_6		= anx2_hed_6.add_run(anx2_hed_str_6)
		anx2_run_hed_7		= anx2_hed_7.add_run(anx2_hed_str_7)

		anx2_run_nme_1		= anx2_par_01.add_run(anx2_run_nme_1)
		anx2_run_nme_2		= anx2_par_02.add_run(anx2_run_nme_2)
		anx2_run_nme_3		= anx2_par_03.add_run(anx2_run_nme_3)
		anx2_run_nme_4		= anx2_par_04.add_run(anx2_run_nme_4)

		anx2_run_par_00		= anx2_par_00.add_run(anx2_str_0)
		anx2_run_par_00		= anx2_par_00.add_run(anx2_par_str_00)
		anx2_run_par_01		= anx2_par_01.add_run(anx2_par_str_01)
		anx2_run_par_02		= anx2_par_02.add_run(anx2_par_str_02)
		anx2_run_par_03		= anx2_par_03.add_run(anx2_par_str_03)
		anx2_run_par_04		= anx2_par_04.add_run(anx2_par_str_04)

		anx2_run_par_11		= anx2_par_11.add_run(anx2_par_str_11)
		anx2_run_par_12		= anx2_par_12.add_run(anx2_par_str_12)
		anx2_run_par_21		= anx2_par_21.add_run(anx2_par_str_21)
		anx2_run_par_22		= anx2_par_22.add_run(anx2_par_str_22)
		anx2_run_par_23		= anx2_par_23.add_run(anx2_par_str_23)
		anx2_run_par_24		= anx2_par_24.add_run(anx2_par_str_24)
		anx2_run_par_25		= anx2_par_25.add_run(anx2_par_str_25)
		anx2_run_par_26		= anx2_par_26.add_run(anx2_par_str_26)
		anx2_run_par_31		= anx2_par_31.add_run(anx2_par_str_31)
		anx2_run_par_32		= anx2_par_32.add_run(anx2_par_str_32)
		anx2_run_par_33		= anx2_par_33.add_run(anx2_par_str_33)
		anx2_run_par_34		= anx2_par_34.add_run(anx2_par_str_34)
		anx2_run_par_35		= anx2_par_35.add_run(anx2_par_str_35)
		anx2_run_par_41		= anx2_par_41.add_run(anx2_par_str_41)
		anx2_run_par_42		= anx2_par_42.add_run(anx2_par_str_42)
		anx2_run_par_43		= anx2_par_43.add_run(anx2_par_str_43)
		anx2_run_par_44		= anx2_par_44.add_run(anx2_par_str_44)
		anx2_run_par_51		= anx2_par_51.add_run(anx2_par_str_51)
		anx2_run_par_52		= anx2_par_52.add_run(anx2_par_str_52)
		anx2_run_par_53		= anx2_par_53.add_run(anx2_par_str_53)
		anx2_run_par_54		= anx2_par_54.add_run(anx2_par_str_54)
		anx2_run_par_55		= anx2_par_55.add_run(anx2_par_str_55)
		anx2_run_par_56		= anx2_par_56.add_run(anx2_par_str_56)
		anx2_run_par_61		= anx2_par_61.add_run(anx2_par_str_61)
		anx2_run_par_62		= anx2_par_62.add_run(anx2_par_str_62)
		anx2_run_par_63		= anx2_par_63.add_run(anx2_par_str_63)
		anx2_run_par_64		= anx2_par_64.add_run(anx2_par_str_64)
		anx2_run_par_71		= anx2_par_71.add_run(anx2_par_str_71)
		anx2_run_par_72		= anx2_par_72.add_run(anx2_par_str_72)

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# M A K E   D O C U M E N T   F O R M A T T I N G
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		run_ttl_1.bold			= False
		run_ttl_2.bold 			= False
		run_ttl_3.bold 			= True
		run_ttl_4.bold 			= False
		run_ttl_5.bold 			= False
		run_ttl_6.bold 			= True
		run_ttl_7.bold 			= False

		#run_hed_3.bold			= False
		run_hed_4.bold			= False
		anx1_run_hed_0.bold		= False
		anx1_run_hed_1.bold		= False
		anx1_run_hed_2.bold		= False
		anx1_run_hed_3.bold		= False
		anx1_run_hed_4.bold		= False
		anx1_run_hed_5.bold		= False
		anx1_run_hed_6.bold		= False
		anx1_run_hed_7.bold		= False
		anx1_run_hed_8.bold		= False

		anx2_run_hed_0.bold		= False
		anx2_run_hed_1.bold		= False
		anx2_run_hed_2.bold		= False
		anx2_run_hed_3.bold		= False
		anx2_run_hed_4.bold		= False
		anx2_run_hed_5.bold		= False
		anx2_run_hed_6.bold		= False
		anx2_run_hed_7.bold		= False

		anx1_run_nme_1.underline	= True
		anx1_run_nme_2.underline	= True
		anx1_run_nme_3.underline	= True
		anx1_run_nme_4.underline	= True
		anx2_run_nme_1.underline	= True
		anx2_run_nme_2.underline	= True
		anx2_run_nme_3.underline	= True
		anx2_run_nme_4.underline	= True

		ttl_emp_fnt_0			= run_emp_0.font
		ttl_emp_fnt_1			= run_emp_1.font
		ttl_emp_fnt_2			= run_emp_2.font
		ttl_emp_fnt_3			= run_emp_3.font
		ttl_emp_fnt_4			= run_emp_4.font
		ttl_emp_fnt_5			= run_emp_5.font
		ttl_emp_fnt_6			= run_emp_6.font

		ttl_par_fnt_1			= run_ttl_1.font
		ttl_par_fnt_2			= run_ttl_2.font
		ttl_par_fnt_3			= run_ttl_3.font
		ttl_par_fnt_4			= run_ttl_4.font
		ttl_par_fnt_5			= run_ttl_5.font
		ttl_par_fnt_6			= run_ttl_6.font
		ttl_par_fnt_7			= run_ttl_7.font

		doc_hed_fnt_1			= run_hed_1.font
		doc_hed_fnt_2			= run_hed_2.font
		#doc_hed_fnt_3			= run_hed_3.font
		doc_hed_fnt_4			= run_hed_4.font

		doc_par_fnt_10			= run_par_10.font
		doc_par_fnt_11			= run_par_11.font
		doc_par_fnt_2			= run_par_2.font
		#doc_par_fnt_3			= run_par_3.font
		doc_par_fnt_41			= run_par_41.font
		doc_par_fnt_42			= run_par_42.font
		doc_par_fnt_43			= run_par_43.font
		doc_par_fnt_44			= run_par_44.font
		doc_par_fnt_45			= run_par_45.font

		anx1_hed_fnt_0			= anx1_run_hed_0.font
		anx1_hed_fnt_1			= anx1_run_hed_1.font
		anx1_hed_fnt_2			= anx1_run_hed_2.font
		anx1_hed_fnt_3			= anx1_run_hed_3.font
		anx1_hed_fnt_4			= anx1_run_hed_4.font
		anx1_hed_fnt_5			= anx1_run_hed_5.font
		anx1_hed_fnt_6			= anx1_run_hed_6.font
		anx1_hed_fnt_7			= anx1_run_hed_7.font
		anx1_hed_fnt_8			= anx1_run_hed_8.font

		anx1_par_fnt_00			= anx1_run_par_00.font
		anx1_par_fnt_01			= anx1_run_par_01.font
		anx1_par_fnt_02			= anx1_run_par_02.font
		anx1_par_fnt_03			= anx1_run_par_03.font
		anx1_par_fnt_04			= anx1_run_par_04.font
		anx1_nme_fnt_01			= anx1_run_nme_1.font
		anx1_nme_fnt_02			= anx1_run_nme_2.font
		anx1_nme_fnt_03			= anx1_run_nme_3.font
		anx1_nme_fnt_04			= anx1_run_nme_4.font
		anx1_par_fnt_11			= anx1_run_par_11.font
		anx1_par_fnt_12			= anx1_run_par_12.font
		anx1_par_fnt_20			= anx1_run_par_20.font
		anx1_par_fnt_21			= anx1_run_par_21.font
		anx1_par_fnt_22			= anx1_run_par_22.font
		anx1_par_fnt_23			= anx1_run_par_23.font
		anx1_par_fnt_24			= anx1_run_par_24.font
		anx1_par_fnt_25			= anx1_run_par_25.font
		anx1_par_fnt_26			= anx1_run_par_26.font
		anx1_par_fnt_27			= anx1_run_par_27.font
		anx1_par_fnt_28			= anx1_run_par_28.font
		anx1_par_fnt_29			= anx1_run_par_29.font
		anx1_par_fnt_30			= anx1_run_par_30.font
		anx1_par_fnt_31			= anx1_run_par_31.font
		anx1_par_fnt_32			= anx1_run_par_32.font
		anx1_par_fnt_33			= anx1_run_par_33.font
		anx1_par_fnt_41			= anx1_run_par_41.font
		anx1_par_fnt_42			= anx1_run_par_42.font
		anx1_par_fnt_43			= anx1_run_par_43.font
		anx1_par_fnt_51			= anx1_run_par_51.font
		anx1_par_fnt_52			= anx1_run_par_52.font
		anx1_par_fnt_53			= anx1_run_par_53.font
		anx1_par_fnt_61			= anx1_run_par_61.font
		anx1_par_fnt_62			= anx1_run_par_62.font
		anx1_par_fnt_63			= anx1_run_par_63.font
		anx1_par_fnt_64			= anx1_run_par_64.font
		anx1_par_fnt_65			= anx1_run_par_65.font
		anx1_par_fnt_71			= anx1_run_par_71.font
		anx1_par_fnt_72			= anx1_run_par_72.font
		anx1_par_fnt_73			= anx1_run_par_73.font
		anx1_par_fnt_81			= anx1_run_par_81.font
		anx1_par_fnt_82			= anx1_run_par_82.font

		anx2_hed_fnt_0			= anx2_run_hed_0.font
		anx2_hed_fnt_1			= anx2_run_hed_1.font
		anx2_hed_fnt_2			= anx2_run_hed_2.font
		anx2_hed_fnt_3			= anx2_run_hed_3.font
		anx2_hed_fnt_4			= anx2_run_hed_4.font
		anx2_hed_fnt_5			= anx2_run_hed_5.font
		anx2_hed_fnt_6			= anx2_run_hed_6.font
		anx2_hed_fnt_7			= anx2_run_hed_7.font

		anx2_par_fnt_00			= anx2_run_par_00.font
		anx2_par_fnt_01			= anx2_run_par_01.font
		anx2_par_fnt_02			= anx2_run_par_02.font
		anx2_par_fnt_03			= anx2_run_par_03.font
		anx2_par_fnt_04			= anx2_run_par_04.font
		anx2_nme_fnt_01			= anx2_run_nme_1.font
		anx2_nme_fnt_02			= anx2_run_nme_2.font
		anx2_nme_fnt_03			= anx2_run_nme_3.font
		anx2_nme_fnt_04			= anx2_run_nme_4.font
		anx2_par_fnt_11			= anx2_run_par_11.font
		anx2_par_fnt_12			= anx2_run_par_12.font
		anx2_par_fnt_21			= anx2_run_par_21.font
		anx2_par_fnt_22			= anx2_run_par_22.font
		anx2_par_fnt_23			= anx2_run_par_23.font
		anx2_par_fnt_24			= anx2_run_par_24.font
		anx2_par_fnt_25			= anx2_run_par_25.font
		anx2_par_fnt_26			= anx2_run_par_26.font
		anx2_par_fnt_31			= anx2_run_par_31.font
		anx2_par_fnt_32			= anx2_run_par_32.font
		anx2_par_fnt_33			= anx2_run_par_33.font
		anx2_par_fnt_34			= anx2_run_par_34.font
		anx2_par_fnt_35			= anx2_run_par_35.font
		anx2_par_fnt_41			= anx2_run_par_41.font
		anx2_par_fnt_42			= anx2_run_par_42.font
		anx2_par_fnt_43			= anx2_run_par_43.font
		anx2_par_fnt_44			= anx2_run_par_44.font
		anx2_par_fnt_51			= anx2_run_par_51.font
		anx2_par_fnt_52			= anx2_run_par_52.font
		anx2_par_fnt_53			= anx2_run_par_53.font
		anx2_par_fnt_54			= anx2_run_par_54.font
		anx2_par_fnt_55			= anx2_run_par_55.font
		anx2_par_fnt_56			= anx2_run_par_56.font
		anx2_par_fnt_61			= anx2_run_par_61.font
		anx2_par_fnt_62			= anx2_run_par_62.font
		anx2_par_fnt_63			= anx2_run_par_63.font
		anx2_par_fnt_64			= anx2_run_par_64.font
		anx2_par_fnt_71			= anx2_run_par_71.font
		anx2_par_fnt_72			= anx2_run_par_72.font


		ttl_emp_s				= [ ttl_emp_0, ttl_emp_1, ttl_emp_2, ttl_emp_3, ttl_emp_4, ttl_emp_5, ttl_emp_6, 																			]
		doc_hed_s				= [ doc_hed_1, doc_hed_2, doc_hed_4,																														] 
		anx1_hed_s				= [ anx1_hed_0, anx1_hed_1, anx1_hed_2, anx1_hed_3, anx1_hed_4, anx1_hed_5, anx1_hed_6, anx1_hed_7, anx1_hed_8, 											]
		anx2_hed_s				= [ anx2_hed_0, anx2_hed_1, anx2_hed_2, anx2_hed_3, anx2_hed_4, anx2_hed_5, anx2_hed_6, anx2_hed_7,															]
		ttl_par_s				= [ ttl_par_1, ttl_par_1, ttl_par_3, ttl_par_4, ttl_par_5, ttl_par_6, ttl_par_7,																			]
		doc_par_s				= [ doc_par_10, doc_par_11, doc_par_2, doc_par_41, doc_par_42, doc_par_43, doc_par_44, doc_par_45,												]
		anx1_par_s				= [ anx1_par_00, anx1_par_01, anx1_par_02, anx1_par_03, anx1_par_04, anx1_par_11, anx1_par_12, anx1_par_20, anx1_par_21, anx1_par_22, anx1_par_23, 
									anx1_par_24, anx1_par_25, anx1_par_26, anx1_par_27, anx1_par_28, anx1_par_29, anx1_par_30, anx1_par_31, anx1_par_32,anx1_par_33, anx1_par_41, 
									anx1_par_42, anx1_par_43, anx1_par_51, anx1_par_52,anx1_par_53,anx1_par_61,anx1_par_62,anx1_par_63, anx1_par_64, anx1_par_65, anx1_par_71, anx1_par_72, 
									anx1_par_73, anx1_par_81, anx1_par_82, anx1_par_0																										]
		anx2_par_s 				= [ anx2_par_00, anx2_par_01, anx2_par_02, anx2_par_03, anx2_par_04, anx2_par_11, anx2_par_12, anx2_par_21, anx2_par_22, anx2_par_23, anx2_par_24, 
									anx2_par_25, anx2_par_26, anx2_par_31, anx2_par_32, anx2_par_33, anx2_par_34, anx2_par_35, anx2_par_41, anx2_par_42, anx2_par_43, anx2_par_44, 
									anx2_par_51, anx2_par_52, anx2_par_53, anx2_par_54, anx2_par_55, anx2_par_56, anx2_par_61, anx2_par_62, anx2_par_63, anx2_par_64, anx2_par_71, 
									anx2_par_72, anx2_par_0																																	]

		ttl_emp_fnt_s			= [ ttl_emp_fnt_0, ttl_emp_fnt_1, ttl_emp_fnt_2, ttl_emp_fnt_3, ttl_emp_fnt_4, ttl_emp_fnt_5, ttl_emp_fnt_6,  	]
		doc_hed_fnt_s			= [ doc_hed_fnt_1, doc_hed_fnt_2, doc_hed_fnt_4, 																											]
		anx1_hed_fnt_s			= [ anx1_hed_fnt_0, anx1_hed_fnt_1, anx1_hed_fnt_2, anx1_hed_fnt_3, anx1_hed_fnt_4, anx1_hed_fnt_5,	anx1_hed_fnt_6, anx1_hed_fnt_7, anx1_hed_fnt_8, 		]
		anx2_hed_fnt_s			= [ anx2_hed_fnt_0, anx2_hed_fnt_1, anx2_hed_fnt_2, anx2_hed_fnt_3, anx2_hed_fnt_4, anx2_hed_fnt_5,	anx2_hed_fnt_6, anx2_hed_fnt_7							]																													
		ttl_par_fnt_s			= [ ttl_par_fnt_1, ttl_par_fnt_2, ttl_par_fnt_3, ttl_par_fnt_4, ttl_par_fnt_5, ttl_par_fnt_6, ttl_par_fnt_7,												]
		doc_par_fnt_s			= [ doc_par_fnt_10, doc_par_fnt_11, doc_par_fnt_2,  doc_par_fnt_41, doc_par_fnt_42, doc_par_fnt_43, doc_par_fnt_44, doc_par_fnt_45							]
		anx1_par_fnt_s			= [ anx1_par_fnt_00, anx1_par_fnt_01, anx1_par_fnt_02, anx1_par_fnt_03, anx1_par_fnt_04, anx1_nme_fnt_01, anx1_nme_fnt_02, anx1_nme_fnt_03, anx1_nme_fnt_04, 
									anx1_par_fnt_11, anx1_par_fnt_12, anx1_par_fnt_20, anx1_par_fnt_21, anx1_par_fnt_22, anx1_par_fnt_23, anx1_par_fnt_24, anx1_par_fnt_25, anx1_par_fnt_26, 
									anx1_par_fnt_27, anx1_par_fnt_28, anx1_par_fnt_29, anx1_par_fnt_30, anx1_par_fnt_31, anx1_par_fnt_32, anx1_par_fnt_33, anx1_par_fnt_41, anx1_par_fnt_42, 
									anx1_par_fnt_43, anx1_par_fnt_51, anx1_par_fnt_52, anx1_par_fnt_53, anx1_par_fnt_61, anx1_par_fnt_62, anx1_par_fnt_63, anx1_par_fnt_64, anx1_par_fnt_65, 
									anx1_par_fnt_71, anx1_par_fnt_72, anx1_par_fnt_73, anx1_par_fnt_81, anx1_par_fnt_82, 																	]
		anx2_par_fnt_s			= [ anx2_par_fnt_00, anx2_par_fnt_01, anx2_par_fnt_02, anx2_par_fnt_03, anx2_par_fnt_04, anx2_nme_fnt_01, anx2_nme_fnt_02, anx2_nme_fnt_03, anx2_nme_fnt_04, 
									anx2_par_fnt_11, anx2_par_fnt_12, anx2_par_fnt_21, anx2_par_fnt_22, anx2_par_fnt_23, anx2_par_fnt_24, anx2_par_fnt_25, anx2_par_fnt_26, anx2_par_fnt_31, 
									anx2_par_fnt_32, anx2_par_fnt_33, 
	anx2_par_fnt_34, anx2_par_fnt_35, anx2_par_fnt_41, anx2_par_fnt_42, anx2_par_fnt_43, anx2_par_fnt_44, anx2_par_fnt_51, 
									anx2_par_fnt_52, anx2_par_fnt_53, anx2_par_fnt_54, anx2_par_fnt_55, anx2_par_fnt_56, anx2_par_fnt_61, anx2_par_fnt_62, anx2_par_fnt_63, anx2_par_fnt_64, 
									anx2_par_fnt_71, anx2_par_fnt_72, 																														]

		for ttl_emp 	in ttl_emp_s			: ttl_emp.line_spacing						= 1
		for doc_hed 	in doc_hed_s			: doc_hed.line_spacing						= 1
		for anx1_hed 	in anx1_hed_s			: anx1_hed.line_spacing						= 1
		for anx2_hed 	in anx2_hed_s			: anx2_hed.line_spacing						= 1
		for ttl_par 	in ttl_par_s			: ttl_par.line_spacing						= 1
		for doc_par 	in doc_par_s			: doc_par.line_spacing						= 1
		for anx1_par 	in anx1_par_s			: anx1_par.line_spacing						= 1
		for anx2_par 	in anx2_par_s			: anx2_par.line_spacing						= 1

		for ttl_emp 	in ttl_emp_s			: ttl_emp.paragraph_format.line_spacing		= 1
		for doc_hed 	in doc_hed_s			: doc_hed.paragraph_format.line_spacing		= 1
		for anx1_hed 	in anx1_hed_s			: anx1_hed.paragraph_format.line_spacing	= 1
		for anx2_hed 	in anx2_hed_s			: anx2_hed.paragraph_format.line_spacing	= 1
		for ttl_par 	in ttl_par_s			: ttl_par.paragraph_format.line_spacing		= 1
		for doc_par 	in doc_par_s			: doc_par.paragraph_format.line_spacing		= 1
		for anx1_par 	in anx1_par_s			: anx1_par.paragraph_format.line_spacing	= 1
		for anx2_par 	in anx2_par_s			: anx2_par.paragraph_format.line_spacing	= 1

		for ttl_emp 	in ttl_emp_s			: ttl_emp.paragraph_format.space_before		= Pt( 8 )
		for doc_hed 	in doc_hed_s			: doc_hed.paragraph_format.space_before		= Pt( 8 )
		for anx1_hed 	in anx1_hed_s			: anx1_hed.paragraph_format.space_before	= Pt( 8 )
		for anx2_hed 	in anx2_hed_s			: anx2_hed.paragraph_format.space_before	= Pt( 8 )
		for ttl_par 	in ttl_par_s			: ttl_par.paragraph_format.space_before		= Pt( 8 )
		for doc_par 	in doc_par_s			: doc_par.paragraph_format.space_before		= Pt( 8 )
		for anx1_par 	in anx1_par_s			: anx1_par.paragraph_format.space_before	= Pt( 0 )
		for anx2_par 	in anx2_par_s			: anx2_par.paragraph_format.space_before	= Pt( 0 )

		for ttl_emp 	in ttl_emp_s			: ttl_emp.paragraph_format.space_after		= Pt( 8 )
		for doc_hed 	in doc_hed_s			: doc_hed.paragraph_format.space_after		= Pt( 8 )
		for anx1_hed 	in anx1_hed_s			: anx1_hed.paragraph_format.space_after		= Pt( 2 )
		for anx2_hed 	in anx2_hed_s			: anx2_hed.paragraph_format.space_after		= Pt( 2 )
		for ttl_par 	in ttl_par_s			: ttl_par.paragraph_format.space_after		= Pt( 8 )
		for doc_par 	in doc_par_s			: doc_par.paragraph_format.space_after		= Pt( 8 )
		for anx1_par 	in anx1_par_s			: anx1_par.paragraph_format.space_after		= Pt( 0 )
		for anx2_par 	in anx2_par_s			: anx2_par.paragraph_format.space_after		= Pt( 0 )

		for doc_hed_fnt in doc_hed_fnt_s	: doc_hed_fnt.color.rgb						= RGBColor( 0x00 , 0x92 , 0x34 )
		for anx1_hed_fnt in anx1_hed_fnt_s	: anx1_hed_fnt.color.rgb					= RGBColor( 0x00 , 0x92 , 0x34 )
		for anx2_hed_fnt in anx2_hed_fnt_s	: anx2_hed_fnt.color.rgb					= RGBColor( 0x00 , 0x92 , 0x34 )
		for ttl_par_fnt in ttl_par_fnt_s	: ttl_par_fnt.color.rgb						= RGBColor( 0x00 , 0x00 , 0x00 )
		for doc_par_fnt in doc_par_fnt_s	: doc_par_fnt.color.rgb						= RGBColor( 0x00 , 0x00 , 0x00 )
		for anx1_par_fnt in anx1_par_fnt_s	: anx1_par_fnt.color.rgb					= RGBColor( 0x00 , 0x00 , 0x00 )
		for anx2_par_fnt in anx2_par_fnt_s	: anx2_par_fnt.color.rgb					= RGBColor( 0x00 , 0x00 , 0x00 )

		for doc_hed_fnt in doc_hed_fnt_s	: doc_hed_fnt.name							= 'Calibri'
		for anx1_hed_fnt in anx1_hed_fnt_s	: anx1_hed_fnt.name							= 'Calibri'
		for anx2_hed_fnt in anx2_hed_fnt_s	: anx2_hed_fnt.name							= 'Calibri'
		for ttl_par_fnt in ttl_par_fnt_s	: ttl_par_fnt.name							= 'Calibri'
		for doc_par_fnt in doc_par_fnt_s	: doc_par_fnt.name							= 'Calibri'
		for anx1_par_fnt in anx1_par_fnt_s	: anx1_par_fnt.name							= 'Calibri'
		for anx2_par_fnt in anx2_par_fnt_s	: anx2_par_fnt.name							= 'Calibri'

		for doc_hed_fnt in doc_hed_fnt_s	: doc_hed_fnt.size							= Pt( 11 )
		for anx1_hed_fnt in anx1_hed_fnt_s	: anx1_hed_fnt.size							= Pt( 11 )
		for anx2_hed_fnt in anx2_hed_fnt_s	: anx2_hed_fnt.size							= Pt( 11 )
		for ttl_par_fnt in ttl_par_fnt_s	: ttl_par_fnt.size							= Pt( 11 )
		for doc_par_fnt in doc_par_fnt_s	: doc_par_fnt.size							= Pt( 11 )
		for anx1_par_fnt in anx1_par_fnt_s	: anx1_par_fnt.size							= Pt( 11 )
		for anx2_par_fnt in anx2_par_fnt_s	: anx2_par_fnt.size							= Pt( 11 )

		doc_tbl_1.alignment		= WD_TABLE_ALIGNMENT.CENTER
		doc_tbl_2.alignment		= WD_TABLE_ALIGNMENT.CENTER

		ttl_par_1.alignment 	= WD_ALIGN_PARAGRAPH.CENTER
		ttl_par_2.alignment 	= WD_ALIGN_PARAGRAPH.CENTER
		ttl_par_3.alignment 	= WD_ALIGN_PARAGRAPH.LEFT
		ttl_par_4.alignment 	= WD_ALIGN_PARAGRAPH.JUSTIFY
		ttl_par_5.alignment 	= WD_ALIGN_PARAGRAPH.RIGHT
		ttl_par_6.alignment 	= WD_ALIGN_PARAGRAPH.RIGHT
		ttl_par_7.alignment 	= WD_ALIGN_PARAGRAPH.JUSTIFY

		pic_par_1.alignment		= WD_ALIGN_PARAGRAPH.CENTER
		
		for par 		in tbl_1_par_list		: 	par.alignment 		= WD_ALIGN_PARAGRAPH.RIGHT
		for par 		in tbl_2_par_list		: 	par.alignment 		= WD_ALIGN_PARAGRAPH.RIGHT
		for par 		in tbl_1_par_list_cnr	: 	par.alignment 		= WD_ALIGN_PARAGRAPH.CENTER
		
		for doc_par 	in doc_par_s 			: 	doc_par.alignment 	= WD_ALIGN_PARAGRAPH.JUSTIFY
		for anx1_par 	in anx1_par_s			:	anx1_par.alignment 	= WD_ALIGN_PARAGRAPH.JUSTIFY
		for anx2_par 	in anx2_par_s			:	anx2_par.alignment 	= WD_ALIGN_PARAGRAPH.JUSTIFY

		for cell in doc_tbl_1.columns[0].cells : cell.width = Cm(0.8)
		for cell in doc_tbl_1.columns[1].cells : cell.width = Cm(3.0)
		for cell in doc_tbl_1.columns[2].cells : cell.width = Cm(6.4)
		for cell in doc_tbl_1.columns[3].cells : cell.width = Cm(2.7)
		for cell in doc_tbl_1.columns[4].cells : cell.width = Cm(1.1)
		for cell in doc_tbl_1.columns[5].cells : cell.width = Cm(3.2)

		for cell in doc_tbl_2.columns[0].cells : cell.width = Cm(10.2)
		for cell in doc_tbl_2.columns[1].cells : cell.width = Cm(7.0)

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# S E T   D O C U M E N T   S T Y L E
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		for section in document.sections:
			section.top_margin		= Cm( 2.0 )
			section.bottom_margin	= Cm( 2.0 )
			section.left_margin		= Cm( 3.0 )
			section.right_margin	= Cm( 1.5 )

		style						= document.styles['Normal']
		font						= style.font
		font.name					= 'Calibri'
		font.size					= Pt( 11 )

		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -
		# S A V E   D O C U M E N T
		# - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -

		#jobcalc - get_date_string()
		now = datetime.now()
		date = '{y}.{m}.{d}'.format(y = now.year, m = now.month, d = now.day)
		time = '{h}.{m}.{s}'.format(h = now.hour, m = now.minute, s = now.second)	

		document.save( '{path}\\service_ofr_{date}_{time}.docx'.format( date = date, time = time ) )
		
	except Exception as e:
			with open( '{path}\\docx_elog.txt', "a") as log_file:
				log_file.write( '\n\n{inf_a}\n{inf_b}\n\n'.format( inf_a = traceback.format_tb( sys.exc_info()[2] )[0], inf_b = str( sys.exc_info()[1] ) ) )

		
		
		